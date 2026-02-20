import io, os, time, base64, uuid, threading
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

FILE_MAX_AGE = 30 * 60  # 30 minutes

@app.errorhandler(RequestEntityTooLarge)
def request_entity_too_large(error):
    return jsonify({'error': 'File is too large. Maximum size is 100 MB.'}), 413

ALLOWED_EXT = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXT


def store_file(file_id, pdf_bytes, original_name):
    path = os.path.join(UPLOAD_DIR, f'{file_id}.pdf')
    with open(path, 'wb') as f:
        f.write(pdf_bytes)


def load_file(file_id):
    path = os.path.join(UPLOAD_DIR, f'{file_id}.pdf')
    if not os.path.exists(path):
        return None
    with open(path, 'rb') as f:
        return f.read()


def cleanup_old_files():
    now = time.time()
    for fname in os.listdir(UPLOAD_DIR):
        fpath = os.path.join(UPLOAD_DIR, fname)
        if not os.path.isfile(fpath):
            continue
        try:
            if now - os.path.getmtime(fpath) > FILE_MAX_AGE:
                os.remove(fpath)
        except OSError:
            pass


def _cleanup_loop():
    while True:
        time.sleep(300)  # run every 5 minutes
        cleanup_old_files()

_cleanup_thread = threading.Thread(target=_cleanup_loop, daemon=True)
_cleanup_thread.start()

def docx_to_pdf_bytes(docx_bytes):
    doc = DocxDocument(io.BytesIO(docx_bytes))
    buf = io.BytesIO()
    rl_doc = SimpleDocTemplate(buf, pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = 'Heading1' if para.style.name.startswith('Heading') else 'Normal'
            story.append(Paragraph(para.text, styles[style]))
            story.append(Spacer(1, 0.1 * inch))
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                story.append(Paragraph(row_text, styles['Normal']))
                story.append(Spacer(1, 0.05 * inch))
    if not story:
        story.append(Paragraph('(empty document)', styles['Normal']))
    rl_doc.build(story)
    return buf.getvalue()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    if not file or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400

    original_name = secure_filename(file.filename)
    file_id = str(uuid.uuid4())
    ext = original_name.rsplit('.', 1)[1].lower()
    file_bytes = file.read()

    if ext == 'docx':
        pdf_bytes = docx_to_pdf_bytes(file_bytes)
    else:
        pdf_bytes = file_bytes

    try:
        with fitz.open(stream=pdf_bytes, filetype='pdf') as pdf:
            page_count = len(pdf)
            thumbnails = []
            for i in range(page_count):
                page = pdf[i]
                mat = fitz.Matrix(0.8, 0.8)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes('png')
                b64 = base64.b64encode(img_data).decode()
                thumbnails.append(f'data:image/png;base64,{b64}')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    cleanup_old_files()
    store_file(file_id, pdf_bytes, original_name)

    return jsonify({
        'file_id': file_id,
        'original_name': original_name,
        'page_count': page_count,
        'thumbnails': thumbnails
    })

@app.route('/merge', methods=['POST'])
def merge():
    data = request.get_json(silent=True) or {}
    pages = data.get('pages', [])
    if not isinstance(pages, list) or not pages:
        return jsonify({'error': 'No pages selected'}), 400

    merger = fitz.open()
    merged_count = 0
    for item in pages:
        if not isinstance(item, dict):
            continue
        file_id = item.get('file_id')
        page_idx = item.get('page_index')
        if file_id is None or page_idx is None:
            continue
        try:
            page_idx = int(page_idx)
            rotation = int(item.get('rotation', 0))
        except (TypeError, ValueError):
            continue

        pdf_bytes = load_file(file_id)
        if not pdf_bytes:
            continue

        try:
            with fitz.open(stream=pdf_bytes, filetype='pdf') as src:
                if page_idx < 0 or page_idx >= len(src):
                    continue
                merger.insert_pdf(src, from_page=page_idx, to_page=page_idx)
            merger[-1].set_rotation((merger[-1].rotation + rotation) % 360)
            merged_count += 1
        except Exception as e:
            print(f"Error merging {file_id} page {page_idx}: {e}")
            continue

    if merged_count == 0:
        merger.close()
        return jsonify({'error': 'No valid pages to merge'}), 400

    out_buf = io.BytesIO()
    merger.save(out_buf)
    merger.close()
    out_buf.seek(0)
    return send_file(out_buf, mimetype='application/pdf',
                     as_attachment=True, download_name='dopeoffice_merged.pdf')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)